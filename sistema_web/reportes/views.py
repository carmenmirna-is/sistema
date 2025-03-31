from django.shortcuts import render, redirect
from django.http import HttpResponse
from django.db import connection
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from docx import Document
from docx.shared import Inches
from datetime import datetime

def generar_reportes(request):
    """Vista para mostrar la página de generación de reportes"""
    if not request.session.get('usuario_id') or request.session.get('tipo_usuario') != 'encargado':
        return redirect('login')
    
    return render(request, 'generar_reportes.html')

def generar_reporte(request):
    """Vista para generar reportes con filtros"""
    if not request.session.get('usuario_id') or request.session.get('tipo_usuario') != 'encargado':
        return redirect('login')
    
    # Obtener parámetros del formulario
    formato = request.GET.get('formato', 'pdf')
    fecha_inicio = request.GET.get('fecha_inicio', '')
    fecha_fin = request.GET.get('fecha_fin', '')
    estado = request.GET.get('estado', 'todos')
    
    encargado_id = request.session.get('usuario_id')
    
    # Construir la consulta SQL con filtros
    try:
        with connection.cursor() as cur:
            # Obtener el espacio asignado al encargado
            cur.execute("SELECT id, nombre FROM espacio WHERE encargado_id = %s", (encargado_id,))
            espacio = cur.fetchone()
            
            if not espacio:
                return HttpResponse("No tienes un espacio asignado", status=403)
            
            espacio_id, espacio_nombre = espacio
            
            # Construir la consulta base
            query = """
                SELECT s.id, s.nombre_evento, s.fecha, u.nombre AS usuario_nombre, 
                       s.estado, s.motivo_rechazo, s.fecha_creacion
                FROM solicitud s
                JOIN usuario u ON s.usuario_id = u.id
                WHERE s.espacio_id = %s
            """
            params = [espacio_id]
            
            # Aplicar filtros
            if estado != 'todos':
                query += " AND s.estado = %s"
                params.append(estado)
            
            if fecha_inicio and fecha_fin:
                query += " AND s.fecha BETWEEN %s AND %s"
                params.extend([fecha_inicio, fecha_fin])
            
            # Ejecutar la consulta
            cur.execute(query, params)
            solicitudes = cur.fetchall()
            
            # Obtener estadísticas
            consulta_estadisticas = """
                SELECT 
                    COUNT(*) AS total,
                    SUM(CASE WHEN estado = 'aceptada' THEN 1 ELSE 0 END) AS aceptadas,
                    SUM(CASE WHEN estado = 'rechazada' THEN 1 ELSE 0 END) AS rechazadas,
                    SUM(CASE WHEN estado = 'pendiente' THEN 1 ELSE 0 END) AS pendientes
                FROM solicitud
                WHERE espacio_id = %s
            """
            params_est = [espacio_id]
            if fecha_inicio and fecha_fin:
                consulta_estadisticas += " AND fecha BETWEEN %s AND %s"
                params_est.extend([fecha_inicio, fecha_fin])
            
            cur.execute(consulta_estadisticas, params_est)
            estadisticas = cur.fetchone()
            
    except Exception as e:
        return HttpResponse(f"Error al generar el reporte: {e}", status=500)
    
    # Definir título y subtítulo del reporte
    titulo = f"Reporte de Solicitudes - {espacio_nombre}"
    subtitulo = "Todas las solicitudes"
    
    if estado != 'todos':
        subtitulo = f"Solicitudes {estado}s"
    
    if fecha_inicio and fecha_fin:
        subtitulo += f" del {fecha_inicio} al {fecha_fin}"
    
    # Generar el reporte en el formato solicitado
    if formato == 'pdf':
        return generar_pdf(solicitudes, estadisticas, titulo, subtitulo)
    elif formato == 'excel':
        return generar_excel(solicitudes, estadisticas, titulo, subtitulo)
    elif formato == 'word':
        return generar_word(solicitudes, estadisticas, titulo, subtitulo)
    else:
        return HttpResponse("Formato no válido", status=400)

def generar_pdf(solicitudes, estadisticas, titulo, subtitulo):
    """Genera un reporte en formato PDF usando ReportLab con ajuste automático en todas las columnas"""
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="reporte_{datetime.now().strftime("%Y%m%d")}.pdf"'
    
    doc = SimpleDocTemplate(response, pagesize=letter)
    elements = []
    
    # Estilos
    styles = getSampleStyleSheet()
    title_style = styles['Title']
    subtitle_style = styles['Heading2']
    normal_style = styles['Normal']
    
    # Título, subtítulo y fecha de generación
    elements.append(Paragraph(titulo, title_style))
    elements.append(Paragraph(subtitulo, subtitle_style))
    elements.append(Paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}", normal_style))
    elements.append(Paragraph(" ", normal_style))
    
    # Estadísticas
    elements.append(Paragraph("Estadísticas:", subtitle_style))
    datos_estadisticas = [
        ["Total solicitudes", "Aceptadas", "Rechazadas", "Pendientes"],
        [estadisticas[0], estadisticas[1], estadisticas[2], estadisticas[3]]
    ]
    tabla_estadisticas = Table(datos_estadisticas, colWidths=[100, 100, 100, 100])
    tabla_estadisticas.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(tabla_estadisticas)
    elements.append(Paragraph(" ", normal_style))
    
    # Listado de solicitudes
    elements.append(Paragraph("Listado de solicitudes:", subtitle_style))
    
    # Encabezados y datos de la tabla (se usa Paragraph en campos largos para ajuste automático)
    datos_tabla = [
        ["ID", "Evento", "Fecha", "Solicitante", "Estado"]
    ]
    for sol in solicitudes:
        datos_tabla.append([
            sol[0],
            Paragraph(str(sol[1]), normal_style),
            Paragraph(str(sol[2]), normal_style),
            Paragraph(str(sol[3]), normal_style),
            Paragraph(str(sol[4]), normal_style)
        ])
    
    # Anchos de columna para todas las columnas
    tabla = Table(datos_tabla, colWidths=[30, 150, 100, 150, 80])
    tabla.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    elements.append(tabla)
    
    doc.build(elements)
    return response

def generar_excel(solicitudes, estadisticas, titulo, subtitulo):
    """Genera un reporte en formato Excel configurando el ajuste de texto en todas las celdas"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Reporte de Solicitudes"
    
    # Título, subtítulo y fecha de generación
    ws['A1'] = titulo
    ws['A1'].font = Font(size=16, bold=True)
    ws.merge_cells('A1:E1')
    
    ws['A2'] = subtitulo
    ws['A2'].font = Font(size=14, bold=True)
    ws.merge_cells('A2:E2')
    
    ws['A3'] = f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ws.merge_cells('A3:E3')
    
    # Estadísticas
    ws['A5'] = "Estadísticas:"
    ws['A5'].font = Font(size=12, bold=True)
    ws.merge_cells('A5:E5')
    
    headers_est = ["Total solicitudes", "Aceptadas", "Rechazadas", "Pendientes"]
    for col, header in enumerate(headers_est, start=1):
        cell = ws.cell(row=6, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="DDDDDD")
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
    
    for col, value in enumerate(estadisticas, start=1):
        cell = ws.cell(row=7, column=col, value=value)
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
    
    # Listado de solicitudes
    ws['A9'] = "Listado de solicitudes:"
    ws['A9'].font = Font(size=12, bold=True)
    ws.merge_cells('A9:E9')
    
    headers = ["ID", "Evento", "Fecha", "Solicitante", "Estado"]
    for col, header in enumerate(headers, start=1):
        cell = ws.cell(row=10, column=col)
        cell.value = header
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="DDDDDD")
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
    
    # Datos y ajuste de texto para cada celda
    for row_idx, sol in enumerate(solicitudes, start=11):
        for col_idx, value in enumerate(sol[:5], start=1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(wrap_text=True, horizontal="center")
    
    # Ajustar anchos de columna para todas
    ws.column_dimensions['A'].width = 10   # ID
    ws.column_dimensions['B'].width = 35   # Evento
    ws.column_dimensions['C'].width = 20   # Fecha
    ws.column_dimensions['D'].width = 25   # Solicitante
    ws.column_dimensions['E'].width = 20   # Estado
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="reporte_{datetime.now().strftime("%Y%m%d")}.xlsx"'
    wb.save(response)
    return response

def generar_word(solicitudes, estadisticas, titulo, subtitulo):
    """Genera un reporte en formato Word con ajuste manual de anchos en todas las columnas"""
    doc = Document()
    
    # Título, subtítulo y fecha de generación
    doc.add_heading(titulo, 0)
    doc.add_heading(subtitulo, level=2)
    doc.add_paragraph(f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    doc.add_paragraph()
    
    # Estadísticas
    doc.add_heading('Estadísticas', level=2)
    tabla_est = doc.add_table(rows=2, cols=4)
    tabla_est.style = 'Table Grid'
    
    headers_est = ["Total solicitudes", "Aceptadas", "Rechazadas", "Pendientes"]
    for i, header in enumerate(headers_est):
        tabla_est.cell(0, i).text = header
    for i, value in enumerate(estadisticas):
        tabla_est.cell(1, i).text = str(value)
    
    doc.add_paragraph()
    
    # Listado de solicitudes
    doc.add_heading('Listado de solicitudes', level=2)
    tabla_sol = doc.add_table(rows=1, cols=5)
    tabla_sol.style = 'Table Grid'
    
    headers = ["ID", "Evento", "Fecha", "Solicitante", "Estado"]
    for i, header in enumerate(headers):
        tabla_sol.cell(0, i).text = header
    
    for sol in solicitudes:
        row_cells = tabla_sol.add_row().cells
        row_cells[0].text = str(sol[0])
        row_cells[1].text = str(sol[1])
        row_cells[2].text = str(sol[2])
        row_cells[3].text = str(sol[3])
        row_cells[4].text = str(sol[4])
    
    # Establecer anchos fijos para todas las columnas (ajusta los valores si es necesario)
    ancho_columnas = [Inches(0.8), Inches(3.5), Inches(2), Inches(2.5), Inches(1.5)]
    for row in tabla_sol.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = ancho_columnas[idx]
    
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="reporte_{datetime.now().strftime("%Y%m%d")}.docx"'
    doc.save(response)
    return response
